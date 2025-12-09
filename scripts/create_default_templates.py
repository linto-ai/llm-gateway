#!/usr/bin/env python3
"""
Script to generate default DOCX templates for the document generation system.

Creates professional document templates with placeholders for:
- basic-report.docx: General-purpose report template
- meeting-summary.docx: Meeting notes and summary template

Usage:
    python scripts/create_default_templates.py

The templates will be created in templates/default/
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT


def create_styles(doc: Document) -> None:
    """Create custom styles for the document."""
    styles = doc.styles

    # Heading 1 style
    if 'CustomHeading1' not in [s.name for s in styles]:
        h1 = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
        h1.font.size = Pt(18)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        h1.paragraph_format.space_after = Pt(12)
        h1.paragraph_format.space_before = Pt(24)

    # Heading 2 style
    if 'CustomHeading2' not in [s.name for s in styles]:
        h2 = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        h2.paragraph_format.space_after = Pt(8)
        h2.paragraph_format.space_before = Pt(16)

    # Metadata label style
    if 'MetadataLabel' not in [s.name for s in styles]:
        ml = styles.add_style('MetadataLabel', WD_STYLE_TYPE.PARAGRAPH)
        ml.font.size = Pt(10)
        ml.font.bold = True
        ml.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        ml.paragraph_format.space_after = Pt(2)


def create_basic_report_template(output_dir: Path) -> None:
    """
    Create a basic report template with standard placeholders.

    Placeholders:
    - {{service_name}}: Name of the service
    - {{job_date}}: Date of job completion
    - {{generated_at}}: Document generation timestamp
    - {{output}}: Main job result content
    """
    doc = Document()
    create_styles(doc)

    # Set document margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Header section
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "{{service_name}}"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.style.font.size = Pt(14)
    header_para.style.font.bold = True
    header_para.style.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Report")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # Metadata section
    doc.add_paragraph()
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Date row
    meta_table.cell(0, 0).text = "Date:"
    meta_table.cell(0, 0).paragraphs[0].runs[0].bold = True
    meta_table.cell(0, 1).text = "{{job_date}}"

    # Generated row
    meta_table.cell(1, 0).text = "Generated:"
    meta_table.cell(1, 0).paragraphs[0].runs[0].bold = True
    meta_table.cell(1, 1).text = "{{generated_at}}"

    # Set column widths
    for row in meta_table.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(3)

    # Horizontal line
    doc.add_paragraph()
    hr = doc.add_paragraph()
    hr.paragraph_format.space_after = Pt(6)
    hr_run = hr.add_run("_" * 80)
    hr_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # Content section header
    content_header = doc.add_paragraph("Content", style='Heading 1')

    # Main content placeholder
    content_para = doc.add_paragraph()
    content_para.add_run("{{output}}")
    content_para.paragraph_format.space_after = Pt(12)

    # Footer
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.text = "Page "
    # Note: Page numbers require special handling, using placeholder text

    # Save document
    output_path = output_dir / "basic-report.docx"
    doc.save(str(output_path))
    print(f"Created: {output_path}")


def create_meeting_summary_template(output_dir: Path) -> None:
    """
    Create a meeting summary template with comprehensive placeholders.

    Placeholders:
    - {{title}}: Meeting title (extracted)
    - {{summary}}: Meeting summary (extracted)
    - {{participants}}: List of participants (extracted)
    - {{topics}}: Topics discussed (extracted)
    - {{action_items}}: Action items (extracted)
    - {{key_points}}: Key points (extracted)
    - {{output}}: Full meeting content
    - {{service_name}}: Name of the service
    - {{job_date}}: Date of job completion
    - {{generated_at}}: Document generation timestamp
    """
    doc = Document()
    create_styles(doc)

    # Set document margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "{{service_name}} - Meeting Summary"
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_para.runs[0] if header_para.runs else header_para.add_run()
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Title section
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("{{title}}")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # Subtitle with date
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Meeting Notes - {{job_date}}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # Metadata box (using table)
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.style = 'Table Grid'

    # Participants
    meta_table.cell(0, 0).text = "Participants"
    meta_table.cell(0, 0).paragraphs[0].runs[0].bold = True
    meta_table.cell(0, 1).text = "{{participants}}"

    # Topics
    meta_table.cell(1, 0).text = "Topics"
    meta_table.cell(1, 0).paragraphs[0].runs[0].bold = True
    meta_table.cell(1, 1).text = "{{topics}}"

    # Set column widths
    for row in meta_table.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(5)

    doc.add_paragraph()

    # Summary section
    summary_header = doc.add_paragraph("Executive Summary", style='Heading 1')
    summary_content = doc.add_paragraph()
    summary_content.add_run("{{summary}}")
    summary_content.paragraph_format.space_after = Pt(12)

    # Key Points section
    key_points_header = doc.add_paragraph("Key Points", style='Heading 1')
    key_points_content = doc.add_paragraph()
    key_points_content.add_run("{{key_points}}")
    key_points_content.paragraph_format.space_after = Pt(12)

    # Action Items section
    action_header = doc.add_paragraph("Action Items", style='Heading 1')
    action_content = doc.add_paragraph()
    action_content.add_run("{{action_items}}")
    action_content.paragraph_format.space_after = Pt(12)

    # Horizontal line
    hr = doc.add_paragraph()
    hr_run = hr.add_run("_" * 80)
    hr_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # Full Content section
    full_content_header = doc.add_paragraph("Full Meeting Content", style='Heading 1')
    full_content = doc.add_paragraph()
    full_content.add_run("{{output}}")

    # Footer
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_para.text = "Generated: {{generated_at}}"
    footer_run = footer_para.runs[0] if footer_para.runs else footer_para.add_run()
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Save document
    output_path = output_dir / "meeting-summary.docx"
    doc.save(str(output_path))
    print(f"Created: {output_path}")


def main():
    """Generate all default templates."""
    # Determine output directory
    script_dir = Path(__file__).parent
    project_root = script_dir.parent
    output_dir = project_root / "templates" / "default"

    # Create output directory if needed
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"Creating default templates in: {output_dir}")
    print("-" * 50)

    # Create templates
    create_basic_report_template(output_dir)
    create_meeting_summary_template(output_dir)

    print("-" * 50)
    print("Done! Templates created successfully.")

    # List created files
    print("\nCreated files:")
    for f in output_dir.glob("*.docx"):
        print(f"  - {f.name} ({f.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
