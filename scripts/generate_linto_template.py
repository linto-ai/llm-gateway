#!/usr/bin/env python3
"""Generate the linto-report.docx professional template with LINAGORA/LinTO branding."""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


SCRIPT_DIR = Path(__file__).parent
REPO_DIR = SCRIPT_DIR.parent
ASSETS_DIR = REPO_DIR / "templates" / "assets"
LINAGORA_LOGO = ASSETS_DIR / "linagora-logo.png"
LINTO_STUDIO_LOGO = ASSETS_DIR / "linto-studio-logo.png"
OUTPUT_PATH = REPO_DIR / "templates" / "default" / "linto-report.docx"

DARK = RGBColor(0x2D, 0x2D, 0x2D)
SUBTLE = RGBColor(0xAA, 0xAA, 0xAA)
RULE_COLOR = "D0D0D0"


def remove_borders(table):
    """Remove all borders from a table."""
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = table._tbl.makeelement(qn("w:tblPr"), {})
        table._tbl.insert(0, tblPr)
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = borders.makeelement(qn(f"w:{side}"), {
            qn("w:val"): "none", qn("w:sz"): "0",
            qn("w:space"): "0", qn("w:color"): "auto",
        })
        borders.append(el)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblPr.append(borders)


def add_field(run, field_code):
    """Insert a field (PAGE, NUMPAGES, etc.) into a run."""
    for tag, attrs in [
        ("w:fldChar", {qn("w:fldCharType"): "begin"}),
        ("w:instrText", {}),
        ("w:fldChar", {qn("w:fldCharType"): "end"}),
    ]:
        el = run._r.makeelement(qn(tag), attrs)
        if tag == "w:instrText":
            el.text = f" {field_code} "
        run._r.append(el)


def add_page_x_of_y(paragraph):
    """Add 'X / Y' page numbering."""
    for i, field in enumerate(["PAGE", "NUMPAGES"]):
        if i > 0:
            r = paragraph.add_run(" / ")
            r.font.name = "Calibri Light"
            r.font.size = Pt(7)
            r.font.color.rgb = SUBTLE
        r = paragraph.add_run()
        r.font.name = "Calibri Light"
        r.font.size = Pt(7)
        r.font.color.rgb = SUBTLE
        add_field(r, field)


def set_bottom_border_on_row(table, row_idx, color=RULE_COLOR, size="2"):
    """Set bottom border on all cells of a table row."""
    for cell in table.rows[row_idx].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.makeelement(qn("w:tcBorders"), {})
        bottom = tcBorders.makeelement(qn("w:bottom"), {
            qn("w:val"): "single", qn("w:sz"): size,
            qn("w:space"): "0", qn("w:color"): color,
        })
        tcBorders.append(bottom)
        tcPr.append(tcBorders)


def set_top_border_on_row(table, row_idx, color=RULE_COLOR, size="2"):
    """Set top border on all cells of a table row."""
    for cell in table.rows[row_idx].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.makeelement(qn("w:tcBorders"), {})
        top = tcBorders.makeelement(qn("w:top"), {
            qn("w:val"): "single", qn("w:sz"): size,
            qn("w:space"): "0", qn("w:color"): color,
        })
        tcBorders.append(top)
        tcPr.append(tcBorders)


def setup_header(section):
    """Header: LinTO Studio logo right-aligned, thin line below via table cell border."""
    header = section.header
    header.is_linked_to_previous = False

    # Clear default paragraphs
    for p in header.paragraphs:
        p.clear()

    # Single-cell table to hold the logo with bottom border
    htable = header.add_table(1, 1, width=Cm(16))
    remove_borders(htable)

    cell = htable.cell(0, 0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)

    if LINTO_STUDIO_LOGO.exists():
        r = p.add_run()
        r.add_picture(str(LINTO_STUDIO_LOGO), height=Cm(0.65))

    # Thin bottom border on the cell
    set_bottom_border_on_row(htable, 0)


def setup_footer(section):
    """Footer: thin line above, LINAGORA logo + mention + page X/Y."""
    footer = section.footer
    footer.is_linked_to_previous = False

    # Clear defaults
    for p in footer.paragraphs:
        p.clear()

    # Single-row table with top border for the line
    ftable = footer.add_table(1, 2, width=Cm(16))
    remove_borders(ftable)
    set_top_border_on_row(ftable, 0)

    # Left: LINAGORA logo + short text
    left = ftable.cell(0, 0)
    left.width = Cm(12)
    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lp.paragraph_format.space_before = Pt(4)
    lp.paragraph_format.space_after = Pt(0)

    if LINAGORA_LOGO.exists():
        r = lp.add_run()
        r.add_picture(str(LINAGORA_LOGO), height=Cm(0.25))

    r = lp.add_run("  LinTO — Open source")
    r.font.name = "Calibri Light"
    r.font.size = Pt(6.5)
    r.font.color.rgb = SUBTLE

    # Right: page X / Y
    right = ftable.cell(0, 1)
    right.width = Cm(4)
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(4)
    rp.paragraph_format.space_after = Pt(0)
    add_page_x_of_y(rp)


def generate_template():
    """Generate the linto-report.docx template."""
    doc = Document()

    # Page setup: A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # Styles
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.15

    for level, (name, size, before, after) in enumerate([
        ("Heading 1", 22, 0, 8),
        ("Heading 2", 14, 16, 6),
        ("Heading 3", 12, 12, 4),
    ], 1):
        s = doc.styles[name]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.color.rgb = DARK
        s.font.bold = level > 1
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)

    # Header / Footer
    setup_header(section)
    setup_footer(section)

    # Body
    doc.add_paragraph("{{output}}")

    # Save
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_PATH))
    print(f"Template generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    generate_template()
