#!/usr/bin/env python3
"""
Script to generate a comprehensive DOCX template that demonstrates all extraction placeholders.

This template includes:
- Standard placeholders: job_id, job_date, service_name, flavor_name, organization_name, generated_at
- Extraction placeholders: title, summary, participants, topics, action_items, key_points, date, sentiment
- The main output placeholder

Usage:
    python scripts/create_extraction_test_template.py

The template will be created in templates/default/extraction-test.docx
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT


def create_styles(doc: Document) -> None:
    """Create custom styles for the document."""
    styles = doc.styles

    # Heading 1 style
    if 'CustomHeading1' not in [s.name for s in styles]:
        h1 = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        h1.paragraph_format.space_after = Pt(10)
        h1.paragraph_format.space_before = Pt(18)

    # Heading 2 style
    if 'CustomHeading2' not in [s.name for s in styles]:
        h2 = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
        h2.font.size = Pt(13)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        h2.paragraph_format.space_after = Pt(6)
        h2.paragraph_format.space_before = Pt(12)

    # Metadata label style
    if 'MetadataLabel' not in [s.name for s in styles]:
        ml = styles.add_style('MetadataLabel', WD_STYLE_TYPE.PARAGRAPH)
        ml.font.size = Pt(9)
        ml.font.bold = True
        ml.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Small text style for placeholders demo
    if 'PlaceholderDemo' not in [s.name for s in styles]:
        pd = styles.add_style('PlaceholderDemo', WD_STYLE_TYPE.PARAGRAPH)
        pd.font.size = Pt(9)
        pd.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        pd.font.italic = True


def create_extraction_test_template(output_dir: Path) -> None:
    """
    Create a comprehensive test template demonstrating all extraction placeholders.

    Standard Placeholders:
    - {{job_id}}: Job UUID
    - {{job_date}}: Job completion date
    - {{service_name}}: Service name
    - {{flavor_name}}: Flavor name
    - {{organization_name}}: Organization name
    - {{generated_at}}: Document generation timestamp

    Extraction Placeholders (from metadata):
    - {{title}}: Document/meeting title
    - {{summary}}: Executive summary
    - {{participants}}: List of participants
    - {{topics}}: Topics discussed
    - {{action_items}}: Action items
    - {{key_points}}: Key points
    - {{date}}: Meeting/document date
    - {{sentiment}}: Overall sentiment

    Content Placeholder:
    - {{output}}: Main job result content (markdown-rendered)
    """
    doc = Document()
    create_styles(doc)

    # Set document margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Header section
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "{{service_name}} | {{flavor_name}}"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.runs[0] if header_para.runs else header_para.add_run()
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Main Title (from extraction)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("{{title}}")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    # Subtitle with extracted date
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("{{date}}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # ===================== METADATA SECTION =====================
    meta_header = doc.add_paragraph("Document Information", style='Heading 1')

    # Create metadata table
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = 'Table Grid'

    # Row 1: Job ID
    meta_table.cell(0, 0).text = "Job ID"
    meta_table.cell(0, 0).paragraphs[0].runs[0].bold = True
    meta_table.cell(0, 1).text = "{{job_id}}"

    # Row 2: Job Date
    meta_table.cell(1, 0).text = "Job Date"
    meta_table.cell(1, 0).paragraphs[0].runs[0].bold = True
    meta_table.cell(1, 1).text = "{{job_date}}"

    # Row 3: Organization
    meta_table.cell(2, 0).text = "Organization"
    meta_table.cell(2, 0).paragraphs[0].runs[0].bold = True
    meta_table.cell(2, 1).text = "{{organization_name}}"

    # Row 4: Generated At
    meta_table.cell(3, 0).text = "Generated At"
    meta_table.cell(3, 0).paragraphs[0].runs[0].bold = True
    meta_table.cell(3, 1).text = "{{generated_at}}"

    # Set column widths
    for row in meta_table.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(4.5)

    doc.add_paragraph()

    # ===================== PARTICIPANTS SECTION =====================
    participants_header = doc.add_paragraph("Participants", style='Heading 1')
    participants_content = doc.add_paragraph()
    participants_content.add_run("{{participants}}")
    participants_content.paragraph_format.space_after = Pt(10)

    # ===================== TOPICS SECTION =====================
    topics_header = doc.add_paragraph("Topics Discussed", style='Heading 1')
    topics_content = doc.add_paragraph()
    topics_content.add_run("{{topics}}")
    topics_content.paragraph_format.space_after = Pt(10)

    # ===================== EXECUTIVE SUMMARY SECTION =====================
    summary_header = doc.add_paragraph("Executive Summary", style='Heading 1')
    summary_content = doc.add_paragraph()
    summary_content.add_run("{{summary}}")
    summary_content.paragraph_format.space_after = Pt(10)

    # ===================== KEY POINTS SECTION =====================
    key_points_header = doc.add_paragraph("Key Points", style='Heading 1')
    key_points_content = doc.add_paragraph()
    key_points_content.add_run("{{key_points}}")
    key_points_content.paragraph_format.space_after = Pt(10)

    # ===================== ACTION ITEMS SECTION =====================
    action_header = doc.add_paragraph("Action Items", style='Heading 1')
    action_content = doc.add_paragraph()
    action_content.add_run("{{action_items}}")
    action_content.paragraph_format.space_after = Pt(10)

    # ===================== SENTIMENT SECTION =====================
    sentiment_header = doc.add_paragraph("Overall Sentiment", style='Heading 1')
    sentiment_content = doc.add_paragraph()
    sentiment_content.add_run("{{sentiment}}")
    sentiment_content.paragraph_format.space_after = Pt(10)

    # Horizontal line
    hr = doc.add_paragraph()
    hr_run = hr.add_run("_" * 80)
    hr_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # ===================== FULL CONTENT SECTION =====================
    content_header = doc.add_paragraph("Full Content", style='Heading 1')
    content_para = doc.add_paragraph()
    content_para.add_run("{{output}}")

    # Footer
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.text = "Generated: {{generated_at}} | Service: {{service_name}} | Flavor: {{flavor_name}}"
    footer_run = footer_para.runs[0] if footer_para.runs else footer_para.add_run()
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Save document
    output_path = output_dir / "extraction-test.docx"
    doc.save(str(output_path))
    print(f"Created: {output_path}")
    return output_path


def main():
    """Generate the extraction test template."""
    # Determine output directory
    script_dir = Path(__file__).parent
    project_root = script_dir.parent
    output_dir = project_root / "templates" / "default"

    # Create output directory if needed
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"Creating extraction test template in: {output_dir}")
    print("-" * 50)

    # Create template
    output_path = create_extraction_test_template(output_dir)

    print("-" * 50)
    print("Done! Template created successfully.")
    print(f"\nFile: {output_path.name} ({output_path.stat().st_size} bytes)")

    # List all placeholders
    print("\nPlaceholders included:")
    print("  Standard:")
    print("    - {{job_id}}")
    print("    - {{job_date}}")
    print("    - {{service_name}}")
    print("    - {{flavor_name}}")
    print("    - {{organization_name}}")
    print("    - {{generated_at}}")
    print("  Extraction:")
    print("    - {{title}}")
    print("    - {{summary}}")
    print("    - {{participants}}")
    print("    - {{topics}}")
    print("    - {{action_items}}")
    print("    - {{key_points}}")
    print("    - {{date}}")
    print("    - {{sentiment}}")
    print("  Content:")
    print("    - {{output}}")


if __name__ == "__main__":
    main()
