#!/usr/bin/env python3
"""Document generation service for DOCX/PDF export."""
import json
import logging
import subprocess
import tempfile
from datetime import datetime
from io import BytesIO
from zoneinfo import ZoneInfo
from pathlib import Path
from typing import Optional, Dict, Any

from app.models.document_template import DocumentTemplate
from app.models.job import Job

logger = logging.getLogger(__name__)


class DocumentService:
    """Service for generating documents from templates and job results."""

    TEMPLATES_DIR = Path("/var/www/data/templates")
    # Default template location - used when no template is provided
    DEFAULT_TEMPLATE_PATH = Path(__file__).parent.parent.parent / "templates/default/basic-report.docx"

    # Standard placeholders always available
    STANDARD_PLACEHOLDERS = [
        "output",
        "job_id",
        "job_date",
        "service_name",
        "flavor_name",
        "organization_name",
        "generated_at",
    ]

    @staticmethod
    def _clean_trailing_backslashes(content: str) -> str:
        """
        Remove trailing backslashes at end of lines.

        LLMs sometimes add backslashes for line breaks in markdown,
        but these are not needed when using nl2br extension.
        """
        if not content:
            return content
        import re
        # Remove single backslash at end of lines (but keep escaped backslashes \\)
        return re.sub(r'(?<!\\)\\(?=\n|$)', '', content)

    async def generate_docx(
        self,
        job: Job,
        template: Optional[DocumentTemplate] = None,
        custom_fields: Optional[Dict[str, Any]] = None,
        version_content: Optional[str] = None,
        version_metadata: Optional[Dict[str, Any]] = None,
        timezone: Optional[str] = None,
    ) -> BytesIO:
        """
        Generate DOCX from template with placeholder substitution.

        Args:
            job: Job with result data
            template: Optional document template to use (uses default if not provided)
            custom_fields: Optional additional fields
            version_content: Optional content from a specific version (overrides job.result)
            version_metadata: Optional metadata from a specific version (overrides job.result.extracted_metadata)

        Returns:
            BytesIO containing the generated DOCX
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for document generation")

        # Determine template path
        if template:
            template_path = self.TEMPLATES_DIR / template.file_path
        else:
            # Use built-in default template
            template_path = self.DEFAULT_TEMPLATE_PATH
            if not template_path.exists():
                raise RuntimeError(f"Default template not found at {template_path}")

        doc = Document(template_path)

        placeholders = self.get_placeholders(job, custom_fields, version_content=version_content, version_metadata=version_metadata, timezone=timezone)
        self.substitute_placeholders(doc, placeholders)

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    async def generate_pdf(
        self,
        job: Job,
        template: Optional[DocumentTemplate] = None,
        custom_fields: Optional[Dict[str, Any]] = None,
        version_content: Optional[str] = None,
        version_metadata: Optional[Dict[str, Any]] = None,
        timezone: Optional[str] = None,
    ) -> BytesIO:
        """
        Generate PDF from job result.

        Always uses: DOCX -> LibreOffice -> PDF pipeline.

        Args:
            job: Job with result data
            template: Optional document template (uses default if not provided)
            custom_fields: Optional additional fields
            version_content: Optional content from a specific version (overrides job.result)
            version_metadata: Optional metadata from a specific version (overrides job.result.extracted_metadata)
            timezone: Optional IANA timezone for date formatting (e.g., 'Europe/Paris')

        Returns:
            BytesIO containing the generated PDF
        """
        # Generate DOCX first (handles default template)
        docx_buffer = await self.generate_docx(job, template, custom_fields, version_content=version_content, version_metadata=version_metadata, timezone=timezone)

        # Convert via LibreOffice
        return await self._convert_docx_to_pdf(docx_buffer, job)

    async def generate_html(
        self,
        job: Job,
        template: Optional[DocumentTemplate] = None,
        custom_fields: Optional[Dict[str, Any]] = None,
        version_content: Optional[str] = None,
        version_metadata: Optional[Dict[str, Any]] = None,
        timezone: Optional[str] = None,
    ) -> str:
        """
        Generate HTML preview from job result using mammoth.

        Pipeline: DOCX template → placeholder substitution → mammoth → HTML

        This preserves:
        - Text formatting (bold, italic, headings)
        - Images (embedded as base64)
        - Tables
        - Lists

        Note: Complex layouts (columns, headers/footers) may be simplified.

        Args:
            job: Job with result data
            template: Optional document template (uses default if not provided)
            custom_fields: Optional additional fields
            version_content: Optional content from a specific version (overrides job.result)
            version_metadata: Optional metadata from a specific version (overrides job.result.extracted_metadata)

        Returns:
            HTML string with embedded images
        """
        try:
            import mammoth
        except ImportError:
            raise ImportError("mammoth is required for HTML preview generation")

        # Generate DOCX first (reuse existing logic)
        docx_buffer = await self.generate_docx(job, template, custom_fields, version_content=version_content, version_metadata=version_metadata, timezone=timezone)

        # Convert DOCX to HTML using mammoth
        result = mammoth.convert_to_html(docx_buffer)
        html_content = result.value

        # Log any conversion warnings
        if result.messages:
            for msg in result.messages:
                logger.warning(f"mammoth conversion warning: {msg}")

        # Wrap in a basic HTML structure with styling
        html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }}
        h1, h2, h3, h4, h5, h6 {{
            margin-top: 1.5em;
            margin-bottom: 0.5em;
            color: #222;
        }}
        h1 {{ font-size: 1.8em; border-bottom: 2px solid #eee; padding-bottom: 0.3em; }}
        h2 {{ font-size: 1.5em; border-bottom: 1px solid #eee; padding-bottom: 0.2em; }}
        p {{ margin: 0.8em 0; }}
        ul, ol {{ margin: 0.8em 0; padding-left: 2em; }}
        li {{ margin: 0.3em 0; }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 1em 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 8px 12px;
            text-align: left;
        }}
        th {{
            background-color: #f5f5f5;
            font-weight: 600;
        }}
        tr:nth-child(even) {{
            background-color: #fafafa;
        }}
        img {{
            max-width: 100%;
            height: auto;
        }}
        blockquote {{
            border-left: 4px solid #ddd;
            margin: 1em 0;
            padding-left: 1em;
            color: #666;
        }}
        code {{
            background: #f5f5f5;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: monospace;
        }}
        pre {{
            background: #f5f5f5;
            padding: 1em;
            border-radius: 5px;
            overflow-x: auto;
        }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""

        return html

    async def _convert_docx_to_pdf(self, docx_buffer: BytesIO, job: Job) -> BytesIO:
        """Convert DOCX buffer to PDF using LibreOffice."""
        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = Path(tmpdir) / "document.docx"
            pdf_path = Path(tmpdir) / "document.pdf"

            with open(docx_path, 'wb') as f:
                f.write(docx_buffer.getvalue())

            # LibreOffice conversion
            try:
                subprocess.run(
                    [
                        'libreoffice',
                        '--headless',
                        '--convert-to', 'pdf',
                        '--outdir', tmpdir,
                        str(docx_path)
                    ],
                    check=True,
                    capture_output=True,
                    timeout=60  # 60 second timeout
                )
            except subprocess.CalledProcessError as e:
                logger.error(f"LibreOffice conversion failed: {e.stderr.decode()}")
                raise RuntimeError("PDF conversion failed")
            except FileNotFoundError:
                raise RuntimeError("LibreOffice not installed. PDF export unavailable.")
            except subprocess.TimeoutExpired:
                raise RuntimeError("PDF conversion timed out")

            # Read the generated PDF
            if not pdf_path.exists():
                raise RuntimeError("PDF file was not generated")

            with open(pdf_path, 'rb') as f:
                return BytesIO(f.read())

    def get_placeholders(
        self,
        job: Job,
        custom_fields: Optional[Dict[str, Any]] = None,
        version_content: Optional[str] = None,
        version_metadata: Optional[Dict[str, Any]] = None,
        timezone: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        Get all available placeholders for template substitution.

        Args:
            job: Job with result data
            custom_fields: Optional additional fields
            version_content: Optional content from a specific version (overrides job.result)
            version_metadata: Optional metadata from a specific version (overrides job.result.extracted_metadata)

        Returns:
            Dict mapping placeholder names to values
        """
        # Use version_content if provided, otherwise extract from job result
        output_content = version_content if version_content else self._get_result_content(job)

        tz = None
        if timezone:
            try:
                tz = ZoneInfo(timezone)
            except (KeyError, ValueError):
                logger.warning(f"Invalid timezone '{timezone}', falling back to UTC")
                tz = ZoneInfo("UTC")

        now = datetime.now(tz) if tz else datetime.now()
        job_date = ""
        if job.completed_at:
            completed = job.completed_at.astimezone(tz) if tz else job.completed_at
            job_date = completed.strftime("%Y-%m-%d")

        placeholders = {
            # Standard placeholders
            "output": output_content,
            "job_id": str(job.id),
            "job_date": job_date,
            "service_name": job.service.name if job.service else "",
            "flavor_name": job.flavor.name if job.flavor else "",
            "organization_name": job.organization_id or "",
            "generated_at": now.strftime("%Y-%m-%d %H:%M"),
        }

        # Add extracted metadata fields - use version_metadata if provided, else job.result.extracted_metadata
        extracted_metadata = version_metadata
        if extracted_metadata is None and job.result and isinstance(job.result, dict):
            extracted_metadata = job.result.get('extracted_metadata')

        if extracted_metadata and isinstance(extracted_metadata, dict):
            for key, value in extracted_metadata.items():
                if key.startswith("_"):  # Skip internal fields like _extraction_error
                    continue
                # Don't override standard placeholders with metadata values
                if key in self.STANDARD_PLACEHOLDERS:
                    continue
                if isinstance(value, list):
                    # Format list items appropriately
                    if all(isinstance(v, dict) for v in value):
                        # List of dicts (like action_items with task/assignee)
                        formatted = []
                        for v in value:
                            if isinstance(v, dict):
                                formatted.append(" - ".join(str(x) for x in v.values()))
                            else:
                                formatted.append(str(v))
                        placeholders[key] = "\n".join(formatted)
                    else:
                        placeholders[key] = ", ".join(str(v) for v in value)
                else:
                    placeholders[key] = str(value) if value is not None else ""

        # Add custom fields (override if same name)
        if custom_fields:
            placeholders.update(custom_fields)

        return placeholders

    def substitute_placeholders(self, doc, placeholders: Dict[str, Any]) -> None:
        """
        Replace {{placeholder}} in DOCX paragraphs, tables, headers, footers.
        For {{output}} placeholder, convert markdown to proper DOCX formatting.

        Args:
            doc: python-docx Document object
            placeholders: Dict of placeholder values
        """
        output_content = placeholders.get("output", "")

        def replace_text_simple(text: str) -> str:
            """Replace placeholders except output (handled specially).

            Handles both formats:
            - {{key}} -> simple placeholder
            - {{key: description}} -> placeholder with extraction hint
            """
            import re
            for key, value in placeholders.items():
                if key == "output":
                    continue  # Handle output separately
                # Replace exact match {{key}}
                text = text.replace(f"{{{{{key}}}}}", str(value or ""))
                # Replace {{key: anything}} pattern (placeholder with description)
                pattern = r'\{\{' + re.escape(key) + r':[^}]*\}\}'
                text = re.sub(pattern, str(value or ""), text)
            return text

        def process_paragraph_for_output(para):
            """Check if paragraph contains {{output}} and replace with formatted content."""
            full_text = "".join(run.text for run in para.runs)
            if "{{output}}" in full_text:
                # Clear the paragraph
                for run in para.runs:
                    run.text = ""
                # Insert formatted markdown content after this paragraph
                return True
            return False

        # First pass: replace simple placeholders
        for para in doc.paragraphs:
            for run in para.runs:
                run.text = replace_text_simple(run.text)

        # Second pass: find and replace {{output}} with formatted content
        paragraphs_to_process = []
        for i, para in enumerate(doc.paragraphs):
            full_text = "".join(run.text for run in para.runs)
            if "{{output}}" in full_text:
                paragraphs_to_process.append((i, para))

        # Insert formatted content for each {{output}} placeholder
        for idx, para in paragraphs_to_process:
            # Clear the placeholder
            for run in para.runs:
                run.text = run.text.replace("{{output}}", "")
            # Insert markdown as formatted DOCX content
            self._insert_markdown_content(doc, para, output_content)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = replace_text_simple(run.text)
                            # For tables, just use plain text for output
                            run.text = run.text.replace("{{output}}", output_content)

        # Replace in headers/footers (simple replacement only)
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    for run in para.runs:
                        run.text = replace_text_simple(run.text)
            if section.footer:
                for para in section.footer.paragraphs:
                    for run in para.runs:
                        run.text = replace_text_simple(run.text)

        # Final pass: clean up any remaining unfilled placeholders
        # Matches {{anything}} or {{anything: with hint}}
        import re
        unfilled_pattern = r'\{\{[^}]+\}\}'

        def clean_unfilled(text: str) -> str:
            return re.sub(unfilled_pattern, '', text)

        for para in doc.paragraphs:
            for run in para.runs:
                run.text = clean_unfilled(run.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = clean_unfilled(run.text)

        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    for run in para.runs:
                        run.text = clean_unfilled(run.text)
            if section.footer:
                for para in section.footer.paragraphs:
                    for run in para.runs:
                        run.text = clean_unfilled(run.text)

    def _insert_markdown_content(self, doc, after_para, markdown_content: str) -> None:
        """
        Convert markdown to DOCX formatting and insert after the given paragraph.

        Args:
            doc: python-docx Document object
            after_para: Paragraph to insert content after
            markdown_content: Markdown text to convert
        """
        try:
            from htmldocx import HtmlToDocx
            import markdown
            from copy import deepcopy
        except ImportError:
            # Fallback: just add plain text
            after_para.add_run(markdown_content)
            return

        # Strip wrapping code fences (LLMs often wrap output in ```markdown ... ```)
        stripped = markdown_content.strip()
        if stripped.startswith("```"):
            lines = stripped.split("\n")
            # Remove opening fence (```markdown, ```md, ``` etc.)
            lines = lines[1:]
            # Remove closing fence if present
            if lines and lines[-1].strip() == "```":
                lines = lines[:-1]
            stripped = "\n".join(lines)
        markdown_content = stripped

        # Clean trailing backslashes before conversion
        cleaned_content = self._clean_trailing_backslashes(markdown_content)

        # Convert markdown to HTML
        html_content = markdown.markdown(
            cleaned_content,
            extensions=['tables', 'fenced_code', 'nl2br']
        )

        # Wrap in proper HTML structure
        html_doc = f"<html><body>{html_content}</body></html>"

        # Create a temporary document to convert HTML
        from docx import Document as DocxDocument
        temp_doc = DocxDocument()
        parser = HtmlToDocx()
        parser.add_html_to_document(html_doc, temp_doc)

        # Insert each element from temp_doc after the placeholder paragraph
        para_element = after_para._element
        insert_point = para_element
        inserted_elements = []
        for element in temp_doc.element.body:
            # Skip section properties (sectPr)
            if element.tag.endswith('sectPr'):
                continue
            # Deep copy the element to preserve all attributes and children
            new_element = deepcopy(element)
            insert_point.addnext(new_element)
            insert_point = new_element
            inserted_elements.append(new_element)

        # htmldocx writes hard-coded English styleIds (Heading1, ListBullet, ...).
        # These only resolve in templates built in English Word. For localized
        # templates (e.g. French uses styleId=Titre1) the styleIds don't exist
        # and Word falls back to defaults, losing the template's look and feel.
        # Remap them to the target template's actual styleIds via the canonical
        # <w:name> (which is invariant across languages for built-in styles).
        self._remap_inserted_style_ids(doc, inserted_elements)

    # Canonical <w:name> values used by Word for the built-in styles that
    # htmldocx may emit. Matching is case-insensitive and space-insensitive.
    _HTMLDOCX_STYLE_CANONICAL_NAMES = {
        "Heading1": "heading 1",
        "Heading2": "heading 2",
        "Heading3": "heading 3",
        "Heading4": "heading 4",
        "Heading5": "heading 5",
        "Heading6": "heading 6",
        "Heading7": "heading 7",
        "Heading8": "heading 8",
        "Heading9": "heading 9",
        "Title": "title",
        "Subtitle": "subtitle",
        "ListBullet": "list bullet",
        "ListNumber": "list number",
        "ListParagraph": "list paragraph",
        "Quote": "quote",
        "IntenseQuote": "intense quote",
        "Caption": "caption",
    }

    _STYLE_OVERRIDE_PROPERTY_PREFIX = "style_"

    def _remap_inserted_style_ids(self, doc, inserted_elements) -> None:
        """Rewrite pStyle/@val on freshly inserted paragraphs so they point at
        styles that actually exist in the target template.

        Resolution order per htmldocx styleId (e.g. "Heading1"):
          1. Document custom property `style_heading1` -> explicit override.
          2. Style in the template whose <w:name> matches the canonical name
             ("heading 1") case-insensitively.
          3. Leave the original styleId untouched (legacy behavior).
        """
        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        W = "{" + W_NS + "}"

        def norm(s: str) -> str:
            return (s or "").strip().lower()

        name_to_id: Dict[str, str] = {}
        for style_el in doc.styles.element.findall(W + "style"):
            sid = style_el.get(W + "styleId")
            if not sid:
                continue
            name_el = style_el.find(W + "name")
            if name_el is None:
                continue
            name_val = name_el.get(W + "val")
            if not name_val:
                continue
            # First entry wins so the primary style for a given name beats
            # later duplicates (e.g. linked character styles).
            name_to_id.setdefault(norm(name_val), sid)

        overrides: Dict[str, str] = self._read_style_overrides(doc)

        # Build the final mapping htmldocx_id -> resolved template styleId.
        resolved: Dict[str, str] = {}
        for htmldocx_id, canonical_name in self._HTMLDOCX_STYLE_CANONICAL_NAMES.items():
            override = overrides.get(htmldocx_id.lower())
            if override:
                resolved[htmldocx_id] = override
                continue
            match = name_to_id.get(canonical_name)
            if match and match != htmldocx_id:
                resolved[htmldocx_id] = match

        if not resolved:
            return

        for root in inserted_elements:
            for pstyle in root.iter(W + "pStyle"):
                current = pstyle.get(W + "val")
                target = resolved.get(current)
                if target:
                    pstyle.set(W + "val", target)

    _CUSTOM_PROPS_CONTENT_TYPE = (
        "application/vnd.openxmlformats-officedocument.custom-properties+xml"
    )
    _CUSTOM_PROPS_NS = (
        "http://schemas.openxmlformats.org/officeDocument/2006/custom-properties"
    )
    _CUSTOM_PROPS_VT_NS = (
        "http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes"
    )

    def _read_style_overrides(self, doc) -> Dict[str, str]:
        """Read style overrides from the document's custom properties.

        A template author can declare, via File > Info > Properties > Advanced
        in Word, a custom property like `style_heading1 = TitreOrange` to force
        inserted H1 paragraphs onto that style. Keys are matched
        case-insensitively against htmldocx styleIds (heading1, listbullet...).

        python-docx 1.2 does not expose custom properties, so we read the raw
        `/docProps/custom.xml` part directly.
        """
        import xml.etree.ElementTree as ET

        try:
            package = doc.part.package
        except Exception:
            return {}

        custom_part = None
        try:
            for part in package.iter_parts():
                if getattr(part, "content_type", None) == self._CUSTOM_PROPS_CONTENT_TYPE:
                    custom_part = part
                    break
        except Exception:
            return {}

        if custom_part is None:
            return {}

        try:
            blob = custom_part.blob
            root = ET.fromstring(blob)
        except Exception:
            return {}

        ns = {"p": self._CUSTOM_PROPS_NS, "vt": self._CUSTOM_PROPS_VT_NS}
        overrides: Dict[str, str] = {}
        for prop in root.findall("p:property", ns):
            name = prop.get("name") or ""
            if not name.lower().startswith(self._STYLE_OVERRIDE_PROPERTY_PREFIX):
                continue
            # Value is in a vt:* child element (lpwstr, bstr, i4, ...).
            value_el = None
            for child in prop:
                if child.tag.startswith("{" + self._CUSTOM_PROPS_VT_NS + "}"):
                    value_el = child
                    break
            if value_el is None or not (value_el.text or "").strip():
                continue
            htmldocx_key = name[len(self._STYLE_OVERRIDE_PROPERTY_PREFIX):].lower()
            overrides[htmldocx_key] = value_el.text.strip()
        return overrides

    def _get_result_content(self, job: Job) -> str:
        """Extract text content from job result."""
        if not job.result:
            return ""
        if isinstance(job.result, str):
            return job.result
        elif isinstance(job.result, dict):
            return (
                job.result.get("content")
                or job.result.get("text")
                or job.result.get("output")
                or json.dumps(job.result, indent=2, ensure_ascii=False)
            )
        return str(job.result)

    def get_all_available_placeholders(
        self,
        template: Optional[DocumentTemplate] = None,
        extracted_metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, list]:
        """
        Get all available placeholder categories.

        Returns:
            Dict with 'standard', 'template', and 'metadata' keys
        """
        result = {
            "standard": self.STANDARD_PLACEHOLDERS.copy(),
            "template": [],
            "metadata": [],
        }

        if template and template.placeholders:
            result["template"] = template.placeholders

        if extracted_metadata:
            result["metadata"] = [
                k for k in extracted_metadata.keys()
                if not k.startswith("_")
            ]

        return result


# Singleton instance
document_service = DocumentService()
