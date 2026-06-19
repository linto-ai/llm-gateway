#!/usr/bin/env python3
"""Document template service for DOCX template management."""
import hashlib
import logging
import re
import os
import uuid as uuid_lib
from pathlib import Path
from typing import Optional, List
from uuid import UUID

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, update, or_, and_, func
from fastapi import UploadFile

from app.models.document_template import DocumentTemplate
from app.services.document_service import DocumentService

logger = logging.getLogger(__name__)


class DocumentTemplateService:
    """Service for managing document templates for DOCX/PDF export.

    Templates are scoped hierarchically:
    - System templates: organization_id=NULL, user_id=NULL (visible to all)
    - Organization templates: organization_id=X, user_id=NULL (visible to org X)
    - User templates: organization_id=X, user_id=Y (visible only to user Y)
    """

    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
    ALLOWED_MIME_TYPES = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]
    DOCX_MAGIC_BYTES = b"PK"  # DOCX is a ZIP file

    # Standard placeholders that are system-provided, not extracted.
    # Single source of truth lives in DocumentService.
    STANDARD_PLACEHOLDERS = DocumentService.STANDARD_PLACEHOLDERS

    def __init__(self):
        """Initialize template service."""
        # Use env var for templates dir, default to /var/www/data/templates
        self.TEMPLATES_DIR = Path(os.environ.get("TEMPLATES_DIR", "/var/www/data/templates"))
        # Ensure templates directory exists
        self.TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)

    @staticmethod
    def _normalize_ids(ids: Optional[List[str]]) -> List[str]:
        """Strip, drop blanks/dupes and sort a list of free-form IDs (stable equality)."""
        if not ids:
            return []
        seen = []
        for raw in ids:
            if raw is None:
                continue
            val = str(raw).strip()
            if val and val not in seen:
                seen.append(val)
        return sorted(seen)

    @classmethod
    def _resolve_scope_lists(
        cls,
        allowed_organization_ids: Optional[List[str]],
        allowed_user_ids: Optional[List[str]],
        organization_id_alias: Optional[str] = None,
        user_id_alias: Optional[str] = None,
    ) -> tuple[List[str], List[str]]:
        """Build the effective (orgs, users) access lists.

        The deprecated single organization_id/user_id aliases are folded in only
        when the explicit lists are empty, preserving backward compatibility.
        """
        orgs = cls._normalize_ids(allowed_organization_ids)
        users = cls._normalize_ids(allowed_user_ids)
        if not orgs and organization_id_alias:
            orgs = cls._normalize_ids([organization_id_alias])
        if not users and user_id_alias:
            users = cls._normalize_ids([user_id_alias])
        return orgs, users

    @staticmethod
    def _derive_legacy_scope(
        orgs: List[str], users: List[str]
    ) -> tuple[Optional[str], Optional[str]]:
        """Derive the legacy single organization_id/user_id from the access lists.

        Kept so clients that read a single scope (LinTO Studio) keep working.
        """
        if not orgs and not users:
            return None, None
        if users:
            return (orgs[0] if orgs else None), users[0]
        return orgs[0], None

    async def create_template(
        self,
        db: AsyncSession,
        file: UploadFile,
        name_fr: str,
        name_en: Optional[str] = None,
        description_fr: Optional[str] = None,
        description_en: Optional[str] = None,
        allowed_organization_ids: Optional[List[str]] = None,
        allowed_user_ids: Optional[List[str]] = None,
        organization_id: Optional[str] = None,
        user_id: Optional[str] = None,
        is_default: bool = False,
    ) -> DocumentTemplate:
        """
        Upload and register a new DOCX template.

        Args:
            db: Database session
            file: Uploaded file
            name_fr: French name (required)
            name_en: English name (optional)
            description_fr: French description (optional)
            description_en: English description (optional)
            allowed_organization_ids: Orgs allowed to see the template (empty => system)
            allowed_user_ids: Users allowed to see the template
            organization_id: Deprecated single-org alias (folded into the lists)
            user_id: Deprecated single-user alias (folded into the lists)
            is_default: Whether to set as default

        Returns:
            Created DocumentTemplate

        Raises:
            ValueError: If file is invalid
        """
        orgs, users = self._resolve_scope_lists(
            allowed_organization_ids, allowed_user_ids, organization_id, user_id
        )
        legacy_org, legacy_user = self._derive_legacy_scope(orgs, users)

        # Validate file
        content = await file.read()
        await file.seek(0)

        if len(content) > self.MAX_FILE_SIZE:
            raise ValueError(f"File exceeds maximum size of {self.MAX_FILE_SIZE // (1024*1024)}MB")

        if not content.startswith(self.DOCX_MAGIC_BYTES):
            raise ValueError("Invalid file format. Only DOCX files are accepted.")

        # Calculate file hash (SHA256)
        file_hash = hashlib.sha256(content).hexdigest()

        # Generate unique file path
        template_id = uuid_lib.uuid4()

        # Determine storage directory based on derived scope (cosmetic grouping)
        if legacy_org is None:
            storage_dir = self.TEMPLATES_DIR / "global"
        elif legacy_user is None:
            storage_dir = self.TEMPLATES_DIR / f"org_{legacy_org}"
        else:
            storage_dir = self.TEMPLATES_DIR / f"org_{legacy_org}" / f"user_{legacy_user}"

        storage_dir.mkdir(parents=True, exist_ok=True)

        # Sanitize filename
        safe_filename = self._sanitize_filename(file.filename or "template.docx")
        file_path = storage_dir / f"{template_id}_{safe_filename}"

        # Save file to disk
        with open(file_path, "wb") as f:
            f.write(content)

        # Extract placeholders from the document
        placeholders = self.extract_placeholders(file_path)

        # If setting as default, unset any existing default at the same scope
        if is_default:
            await self._clear_default_for_scope(db, orgs, users)

        # Create database record
        template = DocumentTemplate(
            id=template_id,
            name_fr=name_fr,
            name_en=name_en,
            description_fr=description_fr,
            description_en=description_en,
            allowed_organization_ids=orgs,
            allowed_user_ids=users,
            organization_id=legacy_org,
            user_id=legacy_user,
            file_path=str(file_path.relative_to(self.TEMPLATES_DIR)),
            file_name=safe_filename,
            file_size=len(content),
            file_hash=file_hash,
            placeholders=placeholders,
            is_default=is_default,
        )

        db.add(template)
        await db.flush()
        await db.refresh(template)

        logger.info(f"Created template: {name_fr} ({template.id}), scope={template.scope}")
        return template

    async def get_template(
        self,
        db: AsyncSession,
        template_id: UUID
    ) -> Optional[DocumentTemplate]:
        """Get template by ID."""
        result = await db.execute(
            select(DocumentTemplate).where(DocumentTemplate.id == template_id)
        )
        return result.scalar_one_or_none()

    async def list_templates(
        self,
        db: AsyncSession,
        organization_id: Optional[str] = None,
        user_id: Optional[str] = None,
        include_system: bool = True,
        include_all: bool = False,
    ) -> List[DocumentTemplate]:
        """
        List templates visible to the given org/user.

        Visibility rules:
        - System templates (org=NULL, user=NULL): visible to all if include_system=True
        - Org templates (org=X, user=NULL): visible to org X members
        - User templates (org=X, user=Y): visible only to user Y

        Args:
            db: Database session
            organization_id: Organization ID for filtering
            user_id: User ID for filtering
            include_system: Also include system templates (default: True)
            include_all: Return ALL templates regardless of scope (admin mode)

        Returns:
            List of DocumentTemplate objects
        """
        # Admin mode: return all templates without filtering
        if include_all:
            query = select(DocumentTemplate)
            query = query.order_by(DocumentTemplate.created_at.desc())
            result = await db.execute(query)
            return list(result.scalars().all())

        conditions = []

        # System templates: both access lists empty.
        if include_system:
            conditions.append(self._is_system_template())

        # Org templates: caller org is in allowed_organization_ids.
        if organization_id:
            conditions.append(
                DocumentTemplate.allowed_organization_ids.any(organization_id)
            )

        # User templates: caller user is in allowed_user_ids.
        if user_id:
            conditions.append(
                DocumentTemplate.allowed_user_ids.any(user_id)
            )

        # If no conditions, return empty list
        if not conditions:
            return []

        query = select(DocumentTemplate).where(or_(*conditions))
        query = query.order_by(DocumentTemplate.created_at.desc())
        result = await db.execute(query)
        return list(result.scalars().all())

    @staticmethod
    def _is_system_template():
        """SQLAlchemy condition: template visible to everyone (both lists empty)."""
        return and_(
            func.cardinality(DocumentTemplate.allowed_organization_ids) == 0,
            func.cardinality(DocumentTemplate.allowed_user_ids) == 0,
        )

    async def update_template(
        self,
        db: AsyncSession,
        template_id: UUID,
        file: Optional[UploadFile] = None,
        name_fr: Optional[str] = None,
        name_en: Optional[str] = None,
        description_fr: Optional[str] = None,
        description_en: Optional[str] = None,
        allowed_organization_ids: Optional[List[str]] = None,
        allowed_user_ids: Optional[List[str]] = None,
        is_default: Optional[bool] = None,
    ) -> Optional[DocumentTemplate]:
        """
        Update template metadata, scope and/or file.

        Args:
            db: Database session
            template_id: Template to update
            file: New file (optional)
            name_fr: New French name (optional)
            name_en: New English name (optional)
            description_fr: New French description (optional)
            description_en: New English description (optional)
            allowed_organization_ids: Replace the org access list (optional)
            allowed_user_ids: Replace the user access list (optional)
            is_default: Set as default (optional)

        Returns:
            Updated DocumentTemplate or None if not found
        """
        template = await self.get_template(db, template_id)
        if not template:
            return None

        # Update metadata fields
        if name_fr is not None:
            template.name_fr = name_fr
        if name_en is not None:
            template.name_en = name_en
        if description_fr is not None:
            template.description_fr = description_fr
        if description_en is not None:
            template.description_en = description_en

        # Update scope (access lists) if either list was provided
        if allowed_organization_ids is not None or allowed_user_ids is not None:
            orgs, users = self._resolve_scope_lists(
                allowed_organization_ids
                if allowed_organization_ids is not None
                else template.allowed_organization_ids,
                allowed_user_ids
                if allowed_user_ids is not None
                else template.allowed_user_ids,
            )
            template.allowed_organization_ids = orgs
            template.allowed_user_ids = users
            # Keep legacy single-scope columns in sync for backward compat.
            template.organization_id, template.user_id = self._derive_legacy_scope(orgs, users)

        # Handle file update
        if file:
            content = await file.read()
            await file.seek(0)

            if len(content) > self.MAX_FILE_SIZE:
                raise ValueError(f"File exceeds maximum size of {self.MAX_FILE_SIZE // (1024*1024)}MB")

            if not content.startswith(self.DOCX_MAGIC_BYTES):
                raise ValueError("Invalid file format. Only DOCX files are accepted.")

            # Calculate new file hash
            file_hash = hashlib.sha256(content).hexdigest()

            # Delete old file
            old_file_path = self.TEMPLATES_DIR / template.file_path
            if old_file_path.exists():
                old_file_path.unlink()

            # Save new file to same directory
            storage_dir = old_file_path.parent
            safe_filename = self._sanitize_filename(file.filename or "template.docx")
            new_file_path = storage_dir / f"{template_id}_{safe_filename}"

            with open(new_file_path, "wb") as f:
                f.write(content)

            # Update file info
            template.file_path = str(new_file_path.relative_to(self.TEMPLATES_DIR))
            template.file_name = safe_filename
            template.file_size = len(content)
            template.file_hash = file_hash
            template.placeholders = self.extract_placeholders(new_file_path)

        # Handle is_default change
        if is_default is not None and is_default != template.is_default:
            if is_default:
                # Clear existing default at same scope
                await self._clear_default_for_scope(
                    db, template.allowed_organization_ids, template.allowed_user_ids
                )
            template.is_default = is_default

        await db.flush()
        await db.refresh(template)
        return template

    async def delete_template(
        self,
        db: AsyncSession,
        template_id: UUID
    ) -> bool:
        """
        Delete template (file + DB record).

        Args:
            db: Database session
            template_id: Template to delete

        Returns:
            True if deleted, False if not found
        """
        template = await self.get_template(db, template_id)
        if not template:
            return False

        # Delete file from disk
        file_path = self.TEMPLATES_DIR / template.file_path
        if file_path.exists():
            file_path.unlink()

        # Delete database record
        await db.delete(template)

        logger.info(f"Deleted template: {template.name_fr} ({template_id})")
        return True

    async def get_default_template(
        self,
        db: AsyncSession,
        organization_id: Optional[str] = None,
        user_id: Optional[str] = None,
    ) -> Optional[DocumentTemplate]:
        """
        Get the default template for a given scope.

        Searches in order: user default -> org default -> system default.

        Args:
            db: Database session
            organization_id: Organization scope
            user_id: User scope

        Returns:
            Default DocumentTemplate or None
        """
        # First try user-level default
        if user_id:
            result = await db.execute(
                select(DocumentTemplate).where(
                    DocumentTemplate.allowed_user_ids.any(user_id),
                    DocumentTemplate.is_default.is_(True)
                )
            )
            template = result.scalars().first()
            if template:
                return template

        # Then try org-level default
        if organization_id:
            result = await db.execute(
                select(DocumentTemplate).where(
                    DocumentTemplate.allowed_organization_ids.any(organization_id),
                    DocumentTemplate.is_default.is_(True)
                )
            )
            template = result.scalars().first()
            if template:
                return template

        # Finally try system-level default (both access lists empty)
        result = await db.execute(
            select(DocumentTemplate).where(
                self._is_system_template(),
                DocumentTemplate.is_default.is_(True)
            )
        )
        return result.scalars().first()

    def extract_placeholders(self, file_path: Path) -> List[str]:
        """
        Parse DOCX and extract {{placeholder}} patterns.

        Args:
            file_path: Path to DOCX file

        Returns:
            List of placeholder names (without braces)
        """
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx not installed, cannot extract placeholders", exc_info=True)
            return []

        placeholders = set()
        pattern = r'\{\{([^}]+)\}\}'

        try:
            doc = Document(file_path)

            # Search paragraphs
            for para in doc.paragraphs:
                placeholders.update(re.findall(pattern, para.text))

            # Search tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        placeholders.update(re.findall(pattern, cell.text))

            # Search headers/footers
            for section in doc.sections:
                if section.header:
                    for para in section.header.paragraphs:
                        placeholders.update(re.findall(pattern, para.text))
                if section.footer:
                    for para in section.footer.paragraphs:
                        placeholders.update(re.findall(pattern, para.text))

        except Exception as e:
            logger.warning(f"Error extracting placeholders from {file_path}: {e}", exc_info=True)

        return sorted(list(placeholders))

    def parse_placeholder_info(self, placeholder: str) -> dict:
        """
        Parse a placeholder string to extract name and description.

        Handles formats:
        - "name" -> {"name": "name", "description": None}
        - "name: description" -> {"name": "name", "description": "description"}

        Args:
            placeholder: Raw placeholder string

        Returns:
            Dict with name, description, and is_standard
        """
        parts = placeholder.split(":", 1)
        name = parts[0].strip()
        description = parts[1].strip() if len(parts) > 1 else None

        return {
            "name": name,
            "description": description,
            "is_standard": name in self.STANDARD_PLACEHOLDERS,
        }

    def get_file_path(self, template: DocumentTemplate) -> Path:
        """Get absolute file path for a template."""
        return self.TEMPLATES_DIR / template.file_path

    def calculate_file_hash(self, file_path: Path) -> str:
        """Calculate SHA256 hash of a file."""
        sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                sha256.update(chunk)
        return sha256.hexdigest()

    async def _clear_default_for_scope(
        self,
        db: AsyncSession,
        allowed_organization_ids: List[str],
        allowed_user_ids: List[str],
    ) -> None:
        """Clear is_default flag for all templates sharing the exact same scope.

        Scope identity is the pair of (sorted) access lists, so only one template
        is the default within a given audience.
        """
        orgs = self._normalize_ids(allowed_organization_ids)
        users = self._normalize_ids(allowed_user_ids)
        await db.execute(
            update(DocumentTemplate)
            .where(
                DocumentTemplate.allowed_organization_ids == orgs,
                DocumentTemplate.allowed_user_ids == users,
                DocumentTemplate.is_default.is_(True)
            )
            .values(is_default=False)
        )

    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename to prevent path traversal."""
        # Remove any path components
        filename = os.path.basename(filename)
        # Remove potentially dangerous characters
        filename = re.sub(r'[^\w\.\-]', '_', filename)
        # Ensure it ends with .docx
        if not filename.lower().endswith('.docx'):
            filename += '.docx'
        return filename

    async def import_template(
        self,
        db: AsyncSession,
        template_id: UUID,
        target_organization_id: Optional[str] = None,
        target_user_id: Optional[str] = None,
        new_name_fr: Optional[str] = None,
        new_name_en: Optional[str] = None,
    ) -> DocumentTemplate:
        """
        Import a system template to an organization or user scope (copy).

        Args:
            db: Database session
            template_id: Source template ID (must be system or org template)
            target_organization_id: Target organization scope
            target_user_id: Target user scope (requires organization)
            new_name_fr: Optional new French name
            new_name_en: Optional new English name

        Returns:
            New DocumentTemplate copy

        Raises:
            ValueError: If source template not found or scope validation fails
        """
        import shutil

        # Validate scope
        if target_user_id is not None and target_organization_id is None:
            raise ValueError("target_user_id requires target_organization_id")

        # Get source template
        source = await self.get_template(db, template_id)
        if not source:
            raise ValueError("Source template not found")

        # Validate source can be imported (only from higher scope)
        if target_user_id:
            # Importing to user scope - source must be system or org
            if source.user_id is not None:
                raise ValueError("Can only import from system or organization templates")
        elif target_organization_id:
            # Importing to org scope - source must be system
            if source.organization_id is not None:
                raise ValueError("Can only import from system templates")
        else:
            raise ValueError("No target scope specified")

        # Get source file path
        source_path = self.get_file_path(source)
        if not source_path.exists():
            raise ValueError("Source template file not found on disk")

        # Calculate file hash
        file_hash = self.calculate_file_hash(source_path)

        # Generate new template ID and destination path
        new_id = uuid_lib.uuid4()

        if target_user_id:
            dest_dir = self.TEMPLATES_DIR / f"org_{target_organization_id}" / f"user_{target_user_id}"
        else:
            dest_dir = self.TEMPLATES_DIR / f"org_{target_organization_id}"

        dest_dir.mkdir(parents=True, exist_ok=True)

        safe_filename = self._sanitize_filename(source.file_name)
        dest_path = dest_dir / f"{new_id}_{safe_filename}"

        # Copy file
        shutil.copy2(source_path, dest_path)

        # Get file size
        file_size = dest_path.stat().st_size

        # Create new template record
        orgs, users = self._resolve_scope_lists(
            None, None, target_organization_id, target_user_id
        )
        legacy_org, legacy_user = self._derive_legacy_scope(orgs, users)
        new_template = DocumentTemplate(
            id=new_id,
            name_fr=new_name_fr or source.name_fr,
            name_en=new_name_en or source.name_en,
            description_fr=source.description_fr,
            description_en=source.description_en,
            allowed_organization_ids=orgs,
            allowed_user_ids=users,
            organization_id=legacy_org,
            user_id=legacy_user,
            file_path=str(dest_path.relative_to(self.TEMPLATES_DIR)),
            file_name=safe_filename,
            file_size=file_size,
            file_hash=file_hash,
            placeholders=source.placeholders,
            is_default=False,
        )

        db.add(new_template)
        await db.flush()
        await db.refresh(new_template)

        logger.info(
            f"Imported template '{source.name_fr}' ({template_id}) "
            f"to {new_template.scope} scope as '{new_template.name_fr}' ({new_id})"
        )
        return new_template


# Singleton instance
document_template_service = DocumentTemplateService()
