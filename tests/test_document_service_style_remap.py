#!/usr/bin/env python3
"""
Tests for DocumentService._insert_markdown_content style remapping.

The pipeline uses htmldocx which hard-codes English styleIds (Heading1,
ListBullet, ...). Localized templates (e.g. French uses styleId=Titre1 with
w:name="heading 1") must have those styleIds remapped via the canonical name,
otherwise the template's look and feel is lost on the inserted content.
"""
import zipfile
from io import BytesIO
from unittest.mock import MagicMock

import pytest
from docx import Document
from lxml import etree

from app.services.document_service import DocumentService


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = "{" + W_NS + "}"


# =============================================================================
# Helpers
# =============================================================================

def _add_style(doc, style_id: str, name: str, based_on: str | None = None):
    """Inject a minimal <w:style> element into doc.styles.element.

    python-docx's add_style would derive the styleId from the name, which is
    exactly what we want to avoid (we need styleId=Titre1 with name='heading 1').
    """
    style_el = etree.SubElement(doc.styles.element, W + "style")
    style_el.set(W + "type", "paragraph")
    style_el.set(W + "styleId", style_id)
    name_el = etree.SubElement(style_el, W + "name")
    name_el.set(W + "val", name)
    if based_on:
        bo = etree.SubElement(style_el, W + "basedOn")
        bo.set(W + "val", based_on)
    return style_el


def _drop_styles_named(doc, names_lower: set[str]):
    """Remove styles from a Document whose <w:name> (case-insensitive) is in
    the given set. Used to undo python-docx's default English heading styles
    so tests can simulate a localized-only template."""
    root = doc.styles.element
    for style_el in list(root.findall(W + "style")):
        name_el = style_el.find(W + "name")
        if name_el is None:
            continue
        val = (name_el.get(W + "val") or "").strip().lower()
        if val in names_lower:
            root.remove(style_el)


def _make_french_styled_doc():
    """Return a fresh Document with French-style ids (Titre1, Listepuces)
    whose canonical w:name matches Word's built-in names. The default English
    heading/list styles are stripped so the template behaves like one authored
    in a French Word."""
    doc = Document()
    _drop_styles_named(
        doc,
        {
            "heading 1", "heading 2", "heading 3",
            "heading 4", "heading 5", "heading 6",
            "heading 7", "heading 8", "heading 9",
            "list bullet", "list number", "list paragraph",
        },
    )
    _add_style(doc, "Titre1", "heading 1")
    _add_style(doc, "Titre2", "heading 2")
    _add_style(doc, "Titre3", "heading 3")
    _add_style(doc, "Listepuces", "List Bullet")
    return doc


def _make_paragraph_with_pstyle(style_val: str):
    """Build a minimal <w:p> element carrying a pStyle reference."""
    p = etree.Element(W + "p")
    ppr = etree.SubElement(p, W + "pPr")
    pstyle = etree.SubElement(ppr, W + "pStyle")
    pstyle.set(W + "val", style_val)
    r = etree.SubElement(p, W + "r")
    t = etree.SubElement(r, W + "t")
    t.text = "dummy"
    return p


def _pstyle_vals(root):
    return [
        pstyle.get(W + "val")
        for pstyle in root.iter(W + "pStyle")
    ]


# =============================================================================
# _remap_inserted_style_ids
# =============================================================================

class TestRemapInsertedStyleIds:
    def test_remaps_english_ids_to_french_via_canonical_name(self):
        """Heading1 -> Titre1 (and friends) when template has FR style ids
        whose w:name matches Word's canonical english names."""
        svc = DocumentService()
        doc = _make_french_styled_doc()

        inserted = [
            _make_paragraph_with_pstyle("Heading1"),
            _make_paragraph_with_pstyle("Heading2"),
            _make_paragraph_with_pstyle("Heading3"),
            _make_paragraph_with_pstyle("ListBullet"),
        ]

        svc._remap_inserted_style_ids(doc, inserted)

        assert _pstyle_vals(inserted[0]) == ["Titre1"]
        assert _pstyle_vals(inserted[1]) == ["Titre2"]
        assert _pstyle_vals(inserted[2]) == ["Titre3"]
        assert _pstyle_vals(inserted[3]) == ["Listepuces"]

    def test_leaves_ids_untouched_when_template_already_has_english_ids(self):
        """No-op when the target template already uses Heading1/Heading2."""
        svc = DocumentService()
        doc = Document()
        # python-docx's default template already defines Heading 1 etc.
        _add_style(doc, "Heading1", "heading 1")
        _add_style(doc, "Heading2", "heading 2")

        inserted = [
            _make_paragraph_with_pstyle("Heading1"),
            _make_paragraph_with_pstyle("Heading2"),
        ]

        svc._remap_inserted_style_ids(doc, inserted)

        assert _pstyle_vals(inserted[0]) == ["Heading1"]
        assert _pstyle_vals(inserted[1]) == ["Heading2"]

    def test_leaves_unresolvable_id_untouched(self):
        """If the template has no matching w:name, the original styleId is
        left in place (legacy fallback, no regression)."""
        svc = DocumentService()
        doc = Document()
        _drop_styles_named(
            doc,
            {"heading 1", "list bullet", "list number", "list paragraph"},
        )
        # Only heading 1 exists in this template — no bullet list equivalent.
        _add_style(doc, "Titre1", "heading 1")

        inserted = [
            _make_paragraph_with_pstyle("Heading1"),
            _make_paragraph_with_pstyle("ListBullet"),
        ]

        svc._remap_inserted_style_ids(doc, inserted)

        assert _pstyle_vals(inserted[0]) == ["Titre1"]
        assert _pstyle_vals(inserted[1]) == ["ListBullet"]  # unchanged

    def test_custom_property_override_wins_over_name_match(self):
        """An explicit `style_heading1 = SomeStyle` custom property must
        override the canonical-name match."""
        svc = DocumentService()
        doc = _make_french_styled_doc()
        _add_style(doc, "TitreOrange", "Titre Orange")

        # Monkey-patch the override reader to simulate a custom property.
        svc._read_style_overrides = lambda _doc: {"heading1": "TitreOrange"}

        inserted = [
            _make_paragraph_with_pstyle("Heading1"),
            _make_paragraph_with_pstyle("Heading2"),
        ]

        svc._remap_inserted_style_ids(doc, inserted)

        assert _pstyle_vals(inserted[0]) == ["TitreOrange"]
        # Heading2 has no override, falls back to canonical-name match.
        assert _pstyle_vals(inserted[1]) == ["Titre2"]

    def test_no_styles_means_no_crash(self):
        """Empty inserted list is a safe no-op."""
        svc = DocumentService()
        doc = Document()
        svc._remap_inserted_style_ids(doc, [])  # should not raise


# =============================================================================
# _read_style_overrides
# =============================================================================

CUSTOM_XML_TEMPLATE = b"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/custom-properties" xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
<property fmtid="{D5CDD505-2E9C-101B-9397-08002B2CF9AE}" pid="2" name="style_heading1"><vt:lpwstr>TitreOrange</vt:lpwstr></property>
<property fmtid="{D5CDD505-2E9C-101B-9397-08002B2CF9AE}" pid="3" name="style_listbullet"><vt:lpwstr>Listepuces</vt:lpwstr></property>
<property fmtid="{D5CDD505-2E9C-101B-9397-08002B2CF9AE}" pid="4" name="Company"><vt:lpwstr>ACME</vt:lpwstr></property>
</Properties>
"""


class _FakePart:
    def __init__(self, content_type: str, blob: bytes):
        self.content_type = content_type
        self.blob = blob


class _FakePackage:
    def __init__(self, parts):
        self._parts = parts

    def iter_parts(self):
        return iter(self._parts)


def _make_doc_with_custom_props(xml_blob: bytes | None):
    """Return a fake 'doc' with just enough structure to satisfy
    _read_style_overrides."""
    parts = []
    if xml_blob is not None:
        parts.append(
            _FakePart(
                DocumentService._CUSTOM_PROPS_CONTENT_TYPE,
                xml_blob,
            )
        )
    # A decoy part to prove the filter works.
    parts.append(_FakePart("application/xml", b"<ignored/>"))

    doc = MagicMock()
    doc.part.package = _FakePackage(parts)
    return doc


class TestReadStyleOverrides:
    def test_parses_style_prefixed_properties_only(self):
        svc = DocumentService()
        doc = _make_doc_with_custom_props(CUSTOM_XML_TEMPLATE)

        overrides = svc._read_style_overrides(doc)

        assert overrides == {
            "heading1": "TitreOrange",
            "listbullet": "Listepuces",
        }
        # Non-style custom property ("Company") is ignored.
        assert "company" not in overrides

    def test_returns_empty_when_no_custom_part(self):
        svc = DocumentService()
        doc = _make_doc_with_custom_props(None)

        assert svc._read_style_overrides(doc) == {}

    def test_returns_empty_on_malformed_xml(self):
        svc = DocumentService()
        doc = _make_doc_with_custom_props(b"not xml at all <<<")

        assert svc._read_style_overrides(doc) == {}


# =============================================================================
# _insert_markdown_content — end-to-end
# =============================================================================

class TestInsertMarkdownContentEndToEnd:
    def test_headings_adopt_french_styleids_after_insertion(self):
        """Full pipeline: markdown -> htmldocx -> remap -> paragraphs in the
        doc carry the template's localized styleIds."""
        pytest.importorskip("htmldocx")
        pytest.importorskip("markdown")

        svc = DocumentService()
        doc = _make_french_styled_doc()
        target = doc.add_paragraph("placeholder")

        md = (
            "# Mon H1\n\n"
            "Un paragraphe.\n\n"
            "## Mon H2\n\n"
            "- item A\n"
            "- item B\n\n"
            "### Mon H3\n\n"
            "Fin.\n"
        )

        svc._insert_markdown_content(doc, target, md)

        # Save -> reload to assert the serialized document is consistent.
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        with zipfile.ZipFile(buf) as z:
            document_xml = z.read("word/document.xml")
        tree = etree.fromstring(document_xml)

        by_text = {}
        for p in tree.iter(W + "p"):
            text = "".join(t.text or "" for t in p.iter(W + "t"))
            pstyle = p.find(W + "pPr/" + W + "pStyle")
            sid = pstyle.get(W + "val") if pstyle is not None else None
            by_text[text.strip()] = sid

        assert by_text.get("Mon H1") == "Titre1"
        assert by_text.get("Mon H2") == "Titre2"
        assert by_text.get("Mon H3") == "Titre3"
        assert by_text.get("item A") == "Listepuces"
        assert by_text.get("item B") == "Listepuces"
